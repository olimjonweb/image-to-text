from flask import Flask, render_template, request, jsonify, send_file
import os
from werkzeug.utils import secure_filename
import pytesseract
from PIL import Image
import pdf2image
import docx
import io

# Tesseract yo'lini ko'rsatish
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_image(image_path):
    try:
        # Rasmni ochish
        image = Image.open(image_path)
        # Matnni ajratib olish
        text = pytesseract.image_to_string(image, lang='uzb+rus+eng')
        if not text.strip():
            return "Rasmdan matn topilmadi. Iltimos, boshqa rasm tanlang."
        return text.strip()
    except Exception as e:
        print(f"Xatolik yuz berdi: {str(e)}")
        if "tesseract is not installed" in str(e):
            return "Tesseract OCR dasturi topilmadi. Iltimos, dasturni o'rnating."
        return f"Xatolik yuz berdi: {str(e)}"

def convert_pdf_to_word(pdf_path):
    try:
        # PDF ni rasmlarga o'tkazish
        images = pdf2image.convert_from_path(pdf_path)
        
        # Yangi Word hujjati yaratish
        doc = docx.Document()
        
        # Har bir sahifadan matnni ajratib olish va Word ga qo'shish
        for image in images:
            text = extract_text_from_image(image)
            doc.add_paragraph(text)
            doc.add_page_break()
        
        # Natijani saqlash
        output_path = pdf_path.rsplit('.', 1)[0] + '.docx'
        doc.save(output_path)
        return output_path
    except Exception as e:
        return str(e)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Fayl tanlanmagan'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Fayl tanlanmagan'}), 400
    
    if file and allowed_file(file.filename):
        try:
            # Yuklanganlar papkasini yaratish
            if not os.path.exists(app.config['UPLOAD_FOLDER']):
                os.makedirs(app.config['UPLOAD_FOLDER'])
            
            # Faylni saqlash
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Rasmdan matnni ajratib olish
            text = extract_text_from_image(filepath)
            
            # Word faylini yaratish
            doc = docx.Document()
            doc.add_paragraph(text)
            
            # Word faylini saqlash
            output_filename = filename.rsplit('.', 1)[0] + '.docx'
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            doc.save(output_path)
            
            # Natijani qaytarish
            return jsonify({
                'success': True,
                'message': 'Matn muvaffaqiyatli ajratib olindi',
                'text': text,
                'download_url': f'/download/{output_filename}'
            })
            
        except Exception as e:
            error_msg = str(e)
            print(f"Xatolik yuz berdi: {error_msg}")
            if "tesseract is not installed" in error_msg:
                return jsonify({'error': "Tesseract OCR dasturi topilmadi. Iltimos, dasturni o'rnating."}), 500
            return jsonify({'error': f'Xatolik yuz berdi: {error_msg}'}), 500
    
    return jsonify({'error': 'Ruxsat etilmagan fayl formati'}), 400

@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_file(
            os.path.join(app.config['UPLOAD_FOLDER'], filename),
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'error': 'Fayl topilmadi'}), 404

if __name__ == '__main__':
    app.run(debug=True) 